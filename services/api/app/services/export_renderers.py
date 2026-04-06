"""Phase 8 — Matplotlib/Plotly PNG, PDF bảng, docx, Excel (data_clean + results_raw)."""

from __future__ import annotations

import json
import logging
from pathlib import Path
from typing import Any

import numpy as np
import pandas as pd

logger = logging.getLogger(__name__)


def _short_value(value: Any, limit: int = 420) -> str:
    if isinstance(value, dict):
        important_keys = [
            "kind",
            "method",
            "decision",
            "p_value",
            "effect_size",
            "alpha",
            "degrees_of_freedom",
        ]
        compact = {k: value.get(k) for k in important_keys if k in value}
        if not compact:
            compact = value
        txt = json.dumps(compact, ensure_ascii=False)
    elif isinstance(value, list):
        if value and isinstance(value[0], dict):
            txt = json.dumps(value[:3], ensure_ascii=False)
            if len(value) > 3:
                txt += f" ... (+{len(value) - 3} items)"
        else:
            txt = json.dumps(value[:20], ensure_ascii=False)
            if len(value) > 20:
                txt += f" ... (+{len(value) - 20} items)"
    else:
        txt = str(value)
    return txt if len(txt) <= limit else f"{txt[:limit]}..."


def _resolve_pdf_fonts() -> tuple[str, str]:
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    candidates = [
        Path("C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/Arial.ttf"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
        Path("/Library/Fonts/Arial Unicode.ttf"),
    ]
    bold_candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf"),
        Path("C:/Windows/Fonts/Arial Bold.ttf"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
        Path("/Library/Fonts/Arial Bold.ttf"),
    ]

    for p in candidates:
        if p.is_file():
            pdfmetrics.registerFont(TTFont("BitlysisSans", str(p)))
            normal = "BitlysisSans"
            break
    else:
        normal = "Helvetica"

    for p in bold_candidates:
        if p.is_file():
            pdfmetrics.registerFont(TTFont("BitlysisSansBold", str(p)))
            bold = "BitlysisSansBold"
            break
    else:
        bold = "Helvetica-Bold"

    return normal, bold


def render_matplotlib_series_png(df: pd.DataFrame, out_path: Path) -> bool:
    """Line chart cột số đầu tiên → PNG (Agg)."""
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    num_cols = df.select_dtypes(include=[np.number]).columns
    if len(num_cols) < 1:
        return False
    col = num_cols[0]
    out_path.parent.mkdir(parents=True, exist_ok=True)
    fig, ax = plt.subplots(figsize=(7.0, 4.0))
    ax.plot(df[col].to_numpy(), marker="o", markersize=2, linewidth=1)
    ax.set_title(f"Chuỗi: {col}")
    ax.set_xlabel("Chỉ số dòng")
    ax.set_ylabel(str(col))
    fig.tight_layout()
    fig.savefig(out_path, dpi=120)
    plt.close(fig)
    return True


def render_plotly_series_png(df: pd.DataFrame, out_path: Path) -> bool:
    """Plotly → PNG (cần kaleido). Lỗi → False, không chặn export."""
    try:
        import plotly.graph_objects as go
    except ImportError:
        return False
    num_cols = df.select_dtypes(include=[np.number]).columns
    if len(num_cols) < 1:
        return False
    col = num_cols[0]
    out_path.parent.mkdir(parents=True, exist_ok=True)
    fig = go.Figure()
    fig.add_trace(
        go.Scatter(
            y=df[col].tolist(),
            mode="lines+markers",
            name=str(col),
            marker=dict(size=4),
        ),
    )
    fig.update_layout(title=f"Plotly: {col}", xaxis_title="Index", yaxis_title=str(col))
    try:
        fig.write_image(str(out_path), width=800, height=500, scale=1)
    except Exception:  # noqa: BLE001
        logger.warning("plotly.write_image failed (kaleido?)")
        return False
    return True


def render_summary_tables_pdf(result_summary: dict[str, Any] | None, out_path: Path) -> None:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    out_path.parent.mkdir(parents=True, exist_ok=True)
    rows: list[list[Any]] = [["Khóa", "Giá trị (rút gọn)"]]
    if result_summary:
        for k, v in list(result_summary.items())[:50]:
            cell = _short_value(v, limit=520)
            rows.append([str(k), cell])
    else:
        rows.append(["result_summary", "(trống)"])

    doc = SimpleDocTemplate(str(out_path), pagesize=A4)
    story: list[Any] = []
    styles = getSampleStyleSheet()
    normal_font, bold_font = _resolve_pdf_fonts()
    title_style = ParagraphStyle(
        "BitlysisTitle",
        parent=styles["Title"],
        fontName=bold_font,
        fontSize=18,
        leading=22,
    )
    cell_style = ParagraphStyle(
        "BitlysisCell",
        parent=styles["BodyText"],
        fontName=normal_font,
        fontSize=8,
        leading=10,
        wordWrap="CJK",
    )
    header_style = ParagraphStyle(
        "BitlysisHeader",
        parent=styles["BodyText"],
        fontName=bold_font,
        fontSize=8,
        textColor=colors.whitesmoke,
    )

    story.append(Paragraph("Bitlysis - Tom tat ket qua (bang)", title_style))
    story.append(Spacer(1, 12))

    table_rows: list[list[Any]] = []
    for i, (k, v) in enumerate(rows):
        if i == 0:
            table_rows.append([
                Paragraph(str(k), header_style),
                Paragraph(str(v), header_style),
            ])
        else:
            table_rows.append([
                Paragraph(str(k), cell_style),
                Paragraph(str(v), cell_style),
            ])

    t = Table(table_rows, colWidths=[125, 370], repeatRows=1)
    t.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                ("FONTNAME", (0, 0), (-1, 0), bold_font),
                ("FONTNAME", (0, 1), (-1, -1), normal_font),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.black),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ],
        ),
    )
    story.append(t)
    doc.build(story)


def render_docx_report(
    *,
    job_id: str,
    original_filename: str,
    columns: list[str],
    result_summary: dict[str, Any] | None,
    out_path: Path,
    template_path: Path | None,
) -> None:
    from docx import Document

    out_path.parent.mkdir(parents=True, exist_ok=True)
    if template_path is not None and template_path.is_file():
        doc = Document(str(template_path))
    else:
        doc = Document()
        doc.add_heading("Bitlysis — Báo cáo xuất", level=0)
    doc.add_paragraph(f"Job ID: {job_id}")
    doc.add_paragraph(f"Tệp gốc: {original_filename}")
    doc.add_paragraph(f"Số cột: {len(columns)}")
    doc.add_heading("Danh sách cột", level=2)
    doc.add_paragraph(", ".join(columns[:80]) or "(không có)")
    if result_summary:
        doc.add_heading("Tóm tắt engine", level=2)
        eng = result_summary.get("engine", "?")
        doc.add_paragraph(f"Engine: {eng}")
    doc.save(str(out_path))


def render_workbook_clean_and_raw(
    df: pd.DataFrame,
    result_summary: dict[str, Any] | None,
    out_path: Path,
    *,
    max_rows: int,
) -> None:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    sample = df.head(max_rows)
    raw_rows: list[list[str]] = []
    if result_summary:
        for k, v in result_summary.items():
            raw_rows.append([str(k), json.dumps(v, ensure_ascii=False)])
    else:
        raw_rows.append(["result_summary", "null"])
    raw_df = pd.DataFrame(raw_rows, columns=["key", "value_json"])
    with pd.ExcelWriter(out_path, engine="openpyxl") as writer:
        sample.to_excel(writer, sheet_name="data_clean", index=False)
        raw_df.to_excel(writer, sheet_name="results_raw", index=False)
